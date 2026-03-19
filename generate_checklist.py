from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HEADER_INFO = [
    ('Проект', 'Dodo Pizza Moscow (https://dodopizza.ru/moscow)'),
    ('Цель тестирования', 'Дымовое тестирование основного пользовательского функционала веб-сервиса'),
    ('Тип тестирования', 'Smoke testing'),
    ('Инструменты', 'Python 3, Selenium WebDriver 4, pytest'),
    ('Окружение', 'Windows / Google Chrome'),
]

TEST_CASES = [
    ('TC-01', 'Открытие главной страницы', 'Smoke', 'Главная страница Dodo Pizza загружается корректно', 'Успешно'),
    ('TC-02', 'Проверка отображения каталога товаров', 'Smoke', 'На странице отображаются карточки товаров и кнопки выбора', 'Успешно'),
    ('TC-03', 'Скролл до блока «Пиццы» по локатору', 'Smoke', 'Страница прокручивается до целевого блока', 'Успешно'),
    ('TC-04', 'Открытие карточки/конструктора товара', 'Smoke', 'После нажатия «Выбрать» открывается окно товара', 'Успешно'),
    ('TC-05', 'Добавление товара в корзину', 'Smoke', 'Товар добавляется, состояние корзины обновляется', 'Успешно'),
    ('TC-06', 'Открытие окна входа и ввод телефона', 'Smoke', 'Открывается форма входа, поле телефона принимает ввод', 'Успешно'),
]

COLUMN_HEADERS = ['ID', 'Наименование теста', 'Тип', 'Ожидаемый результат', 'Результат']
COL_WIDTHS_CM = [1.4, 6.0, 2.0, 6.0, 2.2]
HEADER_BG_COLOR = 'D9D9D9'


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, params in kwargs.items():
        tag = f'w:{edge}'
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), params.get('val', 'single'))
        element.set(qn('w:sz'), str(params.get('sz', 4)))
        element.set(qn('w:color'), params.get('color', '000000'))
        element.set(qn('w:space'), '0')


def set_cell_shading(cell, fill_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def cell_paragraph(cell, text: str, bold=False, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)


def generate_checklist(output_path='checklist_smoke_dodo.docx'):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('ЧЕК-ЛИСТ ДЫМОВОГО ТЕСТИРОВАНИЯ')
    tr.bold = True
    tr.font.name = 'Times New Roman'
    tr.font.size = Pt(15)

    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for key, value in HEADER_INFO:
        row = info_table.add_row()
        cell_paragraph(row.cells[0], key, bold=True, size_pt=11)
        cell_paragraph(row.cells[1], value, size_pt=11)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(COLUMN_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, head in enumerate(COLUMN_HEADERS):
        hdr[i].width = Cm(COL_WIDTHS_CM[i])
        set_cell_shading(hdr[i], HEADER_BG_COLOR)
        cell_paragraph(hdr[i], head, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(hdr[i], top={}, bottom={}, left={}, right={})

    for case in TEST_CASES:
        row = table.add_row().cells
        for i, value in enumerate(case):
            row[i].width = Cm(COL_WIDTHS_CM[i])
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 4) else WD_ALIGN_PARAGRAPH.LEFT
            cell_paragraph(row[i], value, align=align)
            set_cell_border(row[i], top={}, bottom={}, left={}, right={})

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    nr = note.add_run('Примечание: тесты ориентированы на проверку ключевого функционала витрины, выбора товара и корзины.')
    nr.italic = True
    nr.font.name = 'Times New Roman'
    nr.font.size = Pt(10)

    doc.save(output_path)
    print(f'Сохранено: {output_path}')


if __name__ == '__main__':
    generate_checklist()
